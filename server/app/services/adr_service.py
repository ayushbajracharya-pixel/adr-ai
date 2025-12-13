from fastapi import HTTPException, UploadFile
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
import chromadb
import re
from app.utils.text_cleaner import (
    clean_text_from_bullets,
    one_hot_encode_lists_in_dict,
)
import pypdfium2 as pdfium
import docx
from io import BytesIO
from typing import Dict, List, Optional, Any, Union, Set, Tuple
from datetime import datetime
from botocore.exceptions import ClientError

from app.config.settings import settings
from app.models.schemas import QueryIntent
from app.services.retrieval_service import get_hybrid_retriever
from app.services.uploader_service import UploaderService
from app.chains.extraction_chain import ExtractionChain
from app.constants.misc_constants import canonical_headings_map

collection_name = "adr_collection"


class ADRService:
    def __init__(self) -> None:
        """Initialize the ADRService with embeddings, vector store, LLM, and services."""
        self.embeddings = OpenAIEmbeddings(openai_api_key=settings.OPENAI_API_KEY)
        self.client = chromadb.HttpClient(
            host=settings.CHROMADB_HOST, port=settings.CHROMADB_PORT
        )
        self.vector_store = Chroma(
            client=self.client,
            collection_name=collection_name,
            embedding_function=self.embeddings,
        )
        self.llm = ChatOpenAI(
            model=settings.LLM_MODEL_NAME,
            temperature=settings.LLM_TEMPERATURE,
            openai_api_key=settings.OPENAI_API_KEY,
        )
        self.extraction_chain = ExtractionChain()
        self.uploader_service = UploaderService()

    async def process_adr(self, file: UploadFile) -> Tuple[List[str], Dict[str, Any]]:
        """
        Process uploaded ADR and add to knowledge base.

        Args:
            file: The uploaded file to process

        Returns:
            Tuple of (List of document IDs added to the vector store, File information dict)

        Raises:
            HTTPException: If any step of processing fails
        """
        file_name: Optional[str] = file.filename
        if not file_name:
            raise HTTPException(status_code=400, detail="File must have a filename")
        
        content: bytes = await file.read()

        # Step 1: Upload to S3 with specific error handling
        s3_uri = None
        public_url = None
        try:
            s3_object_name = f"adr_uploads/{file_name}"
            with BytesIO(content) as file_obj:
                uploadResponse = self.uploader_service.upload_fileobj(
                    file_obj, s3_object_name
                )
                s3_uri = uploadResponse.get("s3_uri")
                public_url = uploadResponse.get("public_url")
                if not s3_uri:
                    raise IOError("S3 upload returned a malformed response.")
            print(f"Successfully uploaded {file_name} to S3.")
        except Exception as e:
            # Catch any issues with S3 connection, permissions, etc.
            raise HTTPException(
                status_code=500, detail=f"Failed to upload file to S3: {e}"
            )

        # Step 2: Extract text and handle potential errors
        try:
            text_content = clean_text_from_bullets(
                self._extract_text_from_file_content(content, file_name)
            )
        except Exception as e:
            # Catch errors during text extraction from the file content
            # You might want to log this error and potentially delete the S3 object.
            # self.uploader_service.delete_object(s3_object_name)
            raise HTTPException(
                status_code=500, detail=f"Failed to extract text from document: {e}"
            )

        # Step 3: Extract and transform metadata
        try:
            extracted_metadata = self.extraction_chain.invoke_metadata_chain(
                {"text": text_content}
            )
            extracted_metadata["filename"] = file_name
            extracted_metadata["source"] = file_name
            extracted_metadata["s3_uri"] = s3_uri
            extracted_metadata["public_url"] = public_url

            transformed_metadata = one_hot_encode_lists_in_dict(extracted_metadata)

        except Exception as e:
            # Handle potential failures from the LLM extraction chain
            raise HTTPException(
                status_code=500, detail=f"Failed to extract metadata from document: {e}"
            )

        # Step 4: Split and store documents in the vector store
        try:
            # We can clean the text by removing the short info like status, date, etc. which are already present in the metadata
            # so that our chunks contain only meaningful information.

            # Get dynamic separators based on the specific document
            dynamic_separators = self._get_dynamic_separators(text_content)

            # Initialize the splitter with the dynamic separators
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.CHUNK_SIZE,
                chunk_overlap=settings.CHUNK_OVERLAP,
                separators=dynamic_separators,
            )

            document = Document(
                page_content=text_content, metadata=transformed_metadata
            )
            splits = text_splitter.split_documents([document])

            # Ensure all chunks have metadata
            for split in splits:
                split.metadata.update(transformed_metadata)

            doc_ids = self.vector_store.add_documents(splits)

            # Prepare file information for response
            file_info = {
                "object_key": s3_object_name,
                "filename": file_name,
                "size_bytes": len(content),
                "last_modified": datetime.utcnow().isoformat() + "+00:00",
                "permanent_url": public_url
            }

            # Return document IDs and file information
            return doc_ids, file_info

        except Exception as e:
            # Handle failures during vector store indexing
            raise HTTPException(
                status_code=500, detail=f"Failed to index document in vector store: {e}"
            )

    async def query_adr(
        self,
        query: str,
        conversation_history: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """
        Main function to process a user query, extract intent,
        perform hybrid retrieval, and generate a comprehensive response.
        
        Args:
            query: The user's query
            conversation_history: Optional list of previous messages in format
                [{"role": "user", "content": "..."}, {"role": "assistant", "content": "..."}]
        """
        # Step 0: Validate query quality - skip retrieval for low-quality queries
        query_validation = self._validate_query_quality(query, conversation_history)
        if not query_validation["is_searchable"]:
            return self._generate_conversational_response(query, query_validation["reason"], conversation_history)
        
        # Step 1: Extract intent using the LLM chain
        try:
            intent_info = self.extraction_chain.invoke_intent_chain({"query": query})
            print(f"Extracted Intent: {intent_info.model_dump()}")
        except Exception as e:
            # Fallback to simple retrieval if intent extraction fails
            print(f"Intent extraction failed: {e}. Falling back to basic search.")
            intent_info = QueryIntent(
                technologies=[], requirements=[], compliance_needs=[]
            )

        # Step 2: Perform hybrid retrieval based on the extracted intent
        # Pass query for query type classification if not already set
        intent_dict = intent_info.model_dump()
        if not intent_dict.get("query_type"):
            # Query type will be classified in get_hybrid_retriever
            pass
        retriever = get_hybrid_retriever(self.vector_store, intent_dict, query=query)
        retrieved_docs = retriever.invoke(query)

        # If no results, generate a helpful "no results" response
        if not retrieved_docs:
            return self._generate_no_results_response(query, intent_info, conversation_history)

        # Step 3: Create the enhanced prompt using the intent info and conversation history
        prompt_template = self._create_enhanced_prompt(intent_info, conversation_history)

        # Format the retrieved documents for the prompt context
        context_str = self._format_docs(retrieved_docs)

        # Step 4: Generate the final response using the RAG chain
        rag_chain = (
            {"context": RunnablePassthrough(), "question": RunnablePassthrough()}
            | prompt_template
            | self.llm
            | StrOutputParser()
        )

        response_text = rag_chain.invoke({"context": context_str, "question": query})
        
        # Clean up newline characters from the HTML response
        response_text = self._clean_html_response(response_text)

        # Step 5: Process and format the references
        references = self._create_enhanced_references(retrieved_docs)

        # Combine the generated response and references
        return {"query": query, "response": response_text, "references": references}

    async def delete_file(self, object_key: str) -> Dict[str, str]:
        # object_key = "adr_uploads/ADR-0001_ Use Kafka for Messaging in Microservices Architecture.pdf"
        try:
            s3_uri = self.uploader_service.get_s3_uri(object_key)
            self.vector_store.delete(where={"s3_uri": s3_uri})
            print(
                f"✅ Successfully deleted documents for '{object_key}' from ChromaDB."
            )

        except Exception as e:
            # Note: If this fails, the file is already gone from S3.
            # We raise a different error to indicate a partial failure.
            raise HTTPException(
                status_code=500,
                detail=f"Vector store deletion failed: {e}. File was successfully removed from S3, but its index data remains. You may need to manually re-index.",
            )

        try:
            s3_deleted = self.uploader_service.delete_file(object_key)
            if not s3_deleted:
                # The uploader_service handles its own error logging, so we just raise here.
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to delete file from S3 with object key: {object_key}.",
                )
        except ClientError as e:
            # Catch specific Boto3 client errors (e.g., permissions issues, file not found).
            # We re-raise it as an HTTPException for FastAPI to handle.
            raise HTTPException(
                status_code=500, detail=f"S3 deletion failed due to a client error: {e}"
            )
        except Exception as e:
            # Catch any other unexpected errors during the S3 operation.
            raise HTTPException(
                status_code=500,
                detail=f"An unexpected error occurred during S3 deletion: {e}",
            )

        return {"object_key": object_key}

    def _extract_text_from_pdf(self, file_obj: BytesIO) -> str:
        """Extract text from PDF using pypdfium2 from an in-memory object."""
        try:
            # pypdfium2's PdfDocument can directly accept a file-like object (BytesIO)
            pdf = pdfium.PdfDocument(file_obj)
            text_content: List[str] = []

            for page_num in range(len(pdf)):
                page = pdf[page_num]
                textpage = page.get_textpage()
                text: str = textpage.get_text_range()
                text_content.append(text)

                # Clean up
                textpage.close()
                page.close()

            pdf.close()
            return "\n".join(text_content)

        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")

    def _extract_text_from_docx(self, file_obj: BytesIO) -> str:
        """Extract text from Word document from an in-memory object."""
        try:
            # python-docx's Document can also accept a file-like object
            doc = docx.Document(file_obj)
            text_content: List[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            return "\n".join(text_content)

        except Exception as e:
            raise Exception(f"Failed to extract Word document text: {str(e)}")

    def _extract_text_from_file_content(
        self, file_content: bytes, filename: str
    ) -> str:
        """Extract text based on file extension from in-memory content."""
        file_ext: str = filename.lower().split(".")[-1]

        # Create an in-memory file object from the binary content.
        # This is a key step to make file-path-based libraries work.
        file_obj = BytesIO(file_content)

        if file_ext == "pdf":
            return self._extract_text_from_pdf(file_obj)
        elif file_ext in ["docx", "doc"]:
            return self._extract_text_from_docx(file_obj)
        elif file_ext in ["txt", "md"]:
            # For text files, we can just decode the bytes directly.
            return file_content.decode("utf-8")
        else:
            raise Exception(f"Unsupported file format: {file_ext}")

    def _create_enhanced_prompt(
        self,
        intent_info: QueryIntent,
        conversation_history: Optional[List[Dict[str, str]]] = None,
    ) -> ChatPromptTemplate:
        """Creates a contextual prompt template for generation."""
        # This is where you insert your well-crafted prompt instructions
        base_prompt = """
        You are an expert AI assistant helping with technology decisions based on Architecture Decision Records (ADRs) from past projects.

        Your job is to answer questions using only the provided ADR context.
        You must be concise, structured, and professional.

        Context from ADRs:
        {context}
        """

        # Add conversation history if available
        if conversation_history:
            base_prompt += "\n\nPrevious Conversation History:\n"
            # Include last N messages for context
            for msg in conversation_history[-settings.CONVERSATION_HISTORY_LIMIT :]:
                role = "User" if msg.get("role") == "user" else "Assistant"
                content = msg.get("content", "")
                # Truncate very long messages
                if len(content) > settings.MESSAGE_TRUNCATE_LENGTH:
                    content = content[: settings.MESSAGE_TRUNCATE_LENGTH] + "..."
                base_prompt += f"{role}: {content}\n"
            base_prompt += "\n"

        base_prompt += """
        Current User Query: {question}

        ANALYSIS: The user is asking about a new project with these characteristics:
        """

        if intent_info.use_case:
            base_prompt += f"- Use Case: {intent_info.use_case}\n"

        if intent_info.domain:
            base_prompt += f"- Domain/Industry: {intent_info.domain}\n"

        if intent_info.requirements:
            base_prompt += (
                f"- Key Requirements: {', '.join(intent_info.requirements)}\n"
            )

        if intent_info.technologies:
            base_prompt += (
                f"- Relevant Technologies: {', '.join(intent_info.technologies)}\n"
            )

        if intent_info.compliance_needs:
            base_prompt += f"- Compliance Considerations: {', '.join(intent_info.compliance_needs)}\n"

        base_prompt += """

        INSTRUCTIONS:
        Based on the provided context and conversation history, generate a detailed and comprehensive answer to the user's query.
        Consider the conversation history to provide contextually relevant responses that build upon previous discussions.
        
        RESPONSE FORMAT REQUIREMENTS:
        - Format the entire response as a single, valid HTML block. Do not include any text outside of the HTML tags.
        - CRITICAL: Generate compact HTML without newline characters (\n) between tags. Write the HTML as a continuous string without line breaks.
        - Use clear section headings with <h2> or <h3> tags to organize content logically
        - Use bullet points (<ul> and <li>) for lists, recommendations, and key points - avoid long paragraphs when lists are more appropriate
        - Use appropriate emojis where helpful: 🚀 for recommendations/actions, 📌 for important points, ⚠️ for warnings/cautions, ✅ for confirmations/benefits
        - Use <p> tags for paragraphs, <b> or <strong> for emphasis
        - DO NOT include \n characters anywhere in your response - the HTML should be a single continuous string
        - DO NOT add unnecessary line breaks or extra whitespace between HTML elements
        - DO NOT invent information that is not present in the provided context
        - DO NOT repeat the same information multiple times - be concise and avoid redundancy
        - Write in a clear, conversational style similar to ChatGPT - professional but approachable
        - Ensure the response is clean and ready to be rendered in a web browser

        If no relevant ADRs are found in the context, your response should be a simple HTML paragraph:
        <p>Sorry, there has been no such implementations.</p>
        """
        return ChatPromptTemplate.from_template(base_prompt)

    def _format_docs(self, docs: List[Document]) -> str:
        """Formats retrieved documents into a single string for the prompt."""
        formatted: str = ""
        for doc in docs:
            metadata: Dict[str, Any] = doc.metadata
            formatted += f"--- ADR: {metadata.get('title', 'Unknown Title')} ({metadata.get('adr_number', 'Unknown')}) ---\n"
            formatted += f"Source: {metadata.get('source', 'N/A')}\n"
            formatted += f"Content: {doc.page_content}\n\n"
        return formatted

    def _create_enhanced_references(
        self, docs: List[Document]
    ) -> List[Dict[str, str]]:
        """Creates unique, detailed references from retrieved documents."""
        references: List[Dict[str, str]] = []
        seen_files: Set[str] = set()

        # A simple way to get unique documents, more advanced scoring can be added here
        unique_docs: List[Document] = []
        for doc in docs:
            if doc.metadata.get("filename") not in seen_files:
                unique_docs.append(doc)
                seen_files.add(doc.metadata.get("filename"))

        for doc in unique_docs:
            metadata: Dict[str, Any] = doc.metadata
            s3_uri = metadata.get("s3_uri", "Unknown")
            
            # Regenerate public_url from s3_uri to ensure it uses the correct endpoint (LocalStack or AWS)
            # s3_uri format: s3://bucket-name/object-key
            public_url = "Unknown"
            if s3_uri != "Unknown" and s3_uri.startswith("s3://"):
                try:
                    # Extract object_key from s3_uri (everything after s3://bucket-name/)
                    parts = s3_uri.replace("s3://", "").split("/", 1)
                    if len(parts) == 2:
                        object_key = parts[1]
                        # Use uploader_service to generate the correct URL
                        public_url = self.uploader_service._get_public_url(object_key)
                except Exception as e:
                    print(f"Warning: Could not regenerate URL from s3_uri {s3_uri}: {e}")
                    # Fallback to stored public_url if regeneration fails
                    public_url = metadata.get("public_url", "Unknown")
            else:
                # Fallback to stored public_url if s3_uri is invalid
                public_url = metadata.get("public_url", "Unknown")
            
            ref: Dict[str, str] = {
                "filename": metadata.get("filename", "Unknown"),
                "adr_number": metadata.get("adr_number", "Unknown"),
                "title": metadata.get("title", "Unknown Title"),
                "status": metadata.get("status", "Unknown"),
                "author": metadata.get("author", "Unknown"),
                "date": metadata.get("date", "Unknown"),
                "source": metadata.get("source", "N/A"),
                "public_url": public_url,
                "s3_uri": s3_uri,
            }
            references.append(ref)
        return references

    def _validate_query_quality(
        self,
        query: str,
        conversation_history: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """
        Validates if a query is searchable and should trigger ADR retrieval.
        Uses comprehensive heuristics and LLM-based classification.
        
        Returns:
            Dict with 'is_searchable' (bool) and 'reason' (str) if not searchable
        """
        query = query.strip()
        query_lower = query.lower().strip()
        
        # Quick heuristic checks
        if len(query) < 3:
            return {"is_searchable": False, "reason": "query_too_short"}
        
        # Check for common conversational fillers (case-insensitive)
        conversational_fillers = {
            "yeah", "yes", "yep", "yup", "ok", "okay", "okey", "sure", "cool",
            "thanks", "thank you", "thx", "nice", "great", "awesome", "perfect",
            "alright", "alrighty", "got it", "gotcha", "right", "correct",
            "understood", "i see", "i understand", "makes sense", "sounds good",
            "good", "fine", "okay then", "sure thing", "no problem", "np"
        }
        
        if query_lower in conversational_fillers:
            # Special handling for "thank you" - check if there's conversation history
            if query_lower in {"thanks", "thank you", "thx"} and conversation_history:
                # Check if the last assistant message was a substantive response
                for msg in reversed(conversation_history):
                    if msg.get("role") == "assistant":
                        # If assistant provided a substantive response, this is a thank you
                        return {"is_searchable": False, "reason": "thank_you_after_response"}
            return {"is_searchable": False, "reason": "conversational_filler"}
        
        # Check for meta-questions about the conversation/system itself
        meta_question_patterns = [
            r"what (queries|questions|messages) (have|did) (i|you)",
            r"what (did|have) (i|you) (asked|said|queried)",
            r"show (me )?(my|the) (previous|past|earlier) (queries|questions|messages)",
            r"list (my|the) (queries|questions|messages)",
            r"what (have|did) (we|i) (talked|discussed)",
            r"what (is|are) (my|the) (conversation|chat) (history|log)",
            r"how (many|much) (queries|questions) (have|did) (i|you)",
            r"what (can|do) (you|i) (do|help)",
            r"what (are|is) (your|the) (capabilities|features|functions)",
            r"who (are|is) (you|this)",
            r"what (is|are) (this|you)",
            r"help (me|us)",
            r"what (should|can) (i|we) (ask|query)",
            r"how (do|does) (this|it|you) (work|function)",
        ]
        
        for pattern in meta_question_patterns:
            if re.search(pattern, query_lower, re.IGNORECASE):
                return {"is_searchable": False, "reason": "meta_question"}
        
        # Check for greetings and closings
        greeting_patterns = [
            r"^(hi|hello|hey|greetings|good (morning|afternoon|evening))",
            r"^(bye|goodbye|see (you|ya)|farewell|take care)",
        ]
        
        for pattern in greeting_patterns:
            if re.match(pattern, query_lower):
                return {"is_searchable": False, "reason": "greeting_or_closing"}
        
        # Check word count - queries with only 1-2 words are often not searchable
        words = query.split()
        if len(words) <= 1:
            return {"is_searchable": False, "reason": "insufficient_words"}
        
        # LLM-based classification for comprehensive validation
        # This is the primary method for holistic query classification
        try:
            # Build context about conversation history if available
            history_context = ""
            if conversation_history:
                recent_messages = conversation_history[
                    -settings.MAX_RECENT_MESSAGES :
                ]  # Last N messages for context
                history_context = "\n\nRecent conversation context:\n"
                for msg in recent_messages:
                    role = "User" if msg.get("role") == "user" else "Assistant"
                    content = msg.get("content", "")[
                        : settings.MESSAGE_PREVIEW_LENGTH
                    ]  # Truncate long messages
                    history_context += f"{role}: {content}\n"
            
            classification_prompt = f"""
            You are a query classifier for an Architecture Decision Records (ADRs) knowledge base system.
            
            Determine if the following user message is a SEARCHABLE query that should retrieve ADRs from the knowledge base, or if it's a CONVERSATIONAL message that should not trigger ADR retrieval.
            
            User message: "{query}"
            {history_context}
            
            Respond with ONLY one word: "searchable" or "conversational"
            
            Classify as "conversational" (DO NOT retrieve ADRs) if the message is:
            - An acknowledgment, agreement, or confirmation (yeah, ok, thanks, got it, etc.)
            - A greeting or closing (hi, hello, bye, goodbye, etc.)
            - A meta-question about the conversation itself (e.g., "what queries have I asked?", "what did we discuss?", "show my previous questions")
            - A question about the system's capabilities or how to use it (e.g., "what can you do?", "how does this work?", "what should I ask?")
            - A request for help without a specific technical question
            - A simple statement without asking for information
            - Asking about conversation history, previous messages, or chat logs
            - Not requesting information about architecture, technology, or decisions
            - A follow-up acknowledgment to a previous response (e.g., "thanks" after getting an answer)
            
            Classify as "searchable" (SHOULD retrieve ADRs) if the message is:
            - Asking a substantive question about architecture, technology, or design decisions
            - Requesting information, examples, or guidance about technical topics
            - Describing a use case, requirement, or problem that could match ADR content
            - Asking "what", "how", "why", "when", "where" questions about technical topics
            - Requesting comparisons, recommendations, or best practices
            - Any query that could reasonably match content in ADR documents
            
            Examples:
            - "what queries have I asked?" → conversational (meta-question)
            - "what technologies have we used for messaging?" → searchable (technical question)
            - "thanks" → conversational (acknowledgment)
            - "how do we handle authentication?" → searchable (technical question)
            - "what can you do?" → conversational (system capability question)
            - "what database should we use?" → searchable (technical decision question)
            """
            
            classification_chain = ChatPromptTemplate.from_template(classification_prompt) | self.llm | StrOutputParser()
            classification = classification_chain.invoke({}).strip().lower()
            
            if "conversational" in classification:
                return {"is_searchable": False, "reason": "llm_classified_conversational"}
            
        except Exception as e:
            # If LLM classification fails, default to allowing the query (fail open)
            print(f"Query classification failed: {e}. Allowing query to proceed.")
        
        return {"is_searchable": True, "reason": None}

    def _generate_conversational_response(
        self,
        query: str,
        reason: str,
        conversation_history: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """
        Generates a friendly conversational response for non-searchable queries.
        Does not trigger ADR retrieval. Provides context-aware responses when appropriate.
        """
        query_lower = query.lower().strip()
        
        # Special handling for "thank you" after a response - make it warm and friendly
        if reason == "thank_you_after_response" or (reason == "conversational_filler" and query_lower in {"thanks", "thank you", "thx"}):
            if conversation_history:
                # Check if we just provided a substantive response
                for msg in reversed(conversation_history):
                    if msg.get("role") == "assistant":
                        assistant_content = msg.get("content", "")
                        # If the last assistant message was substantive (not just a greeting)
                        if len(assistant_content) > 50 and not assistant_content.startswith("<p>I'm here to help"):
                            return {
                                "query": query,
                                "response": "<p>You're welcome! Happy to help. Feel free to ask if you need anything else about architecture decisions. 🚀</p>",
                                "references": []
                            }
        
        # Map reasons to appropriate responses
        response_templates = {
            "query_too_short": "<p>I'd be happy to help! Could you please provide more details about what you're looking for? 📌</p>",
            "conversational_filler": "<p>Got it! Feel free to ask me anything about architecture decisions or ADRs. 🚀</p>",
            "insufficient_words": "<p>I'd love to help! Could you provide a bit more detail about what you need? 💡</p>",
            "meta_question": "<p>I'm focused on helping you find information about architecture decisions and ADRs. Ask me questions about technologies, design patterns, or architectural choices, and I'll search through the ADR knowledge base for relevant information. 🔍</p>",
            "greeting_or_closing": "<p>Hello! I'm here to help you find information about architecture decisions and ADRs. What would you like to know? 🚀</p>",
            "llm_classified_conversational": "<p>I'm here to help with questions about architecture decisions and ADRs. What would you like to know? 🔍</p>"
        }
        
        # Default response if reason not in templates
        default_response = "<p>I'm here to help! What would you like to know about architecture decisions? 🚀</p>"
        
        response_html = response_templates.get(reason, default_response)
        
        return {
            "query": query,
            "response": response_html,
            "references": []  # No references for conversational responses
        }

    def _clean_html_response(self, html_text: str) -> str:
        """
        Removes newline characters from HTML response to prevent unnecessary whitespace.
        Preserves spaces within text content but removes newlines between tags.
        """
        # Remove all newline characters (\n) from the HTML
        # This creates compact HTML without extra whitespace
        return html_text.replace('\n', '')

    def _generate_no_results_response(
        self,
        query: str,
        intent_info: QueryIntent,
        conversation_history: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """
        Generates a helpful "no results" response with suggestions for reformulating the query.
        Uses LLM to provide actionable guidance to the user.
        """
        # Build context about what was searched
        search_context = f"User Query: {query}\n\n"
        
        if intent_info.technologies:
            search_context += f"Technologies searched: {', '.join(intent_info.technologies)}\n"
        if intent_info.domain:
            search_context += f"Domain/Industry searched: {intent_info.domain}\n"
        if intent_info.requirements:
            search_context += f"Requirements searched: {', '.join(intent_info.requirements)}\n"
        
        search_context += "\nNo matching ADRs were found in the knowledge base."

        # Create a prompt for generating helpful no-results response
        no_results_prompt = f"""
        You are a helpful AI assistant for an Architecture Decision Records (ADRs) knowledge base.
        
        The user searched for information but no relevant ADRs were found.
        
        Search Details:
        {search_context}
        
        Generate a helpful, empathetic response that:
        1. Acknowledges that no relevant ADRs were found
        2. Suggests ways to reformulate the query (e.g., use different keywords, broader terms, related technologies)
        3. Provide actionable next steps
        4. Maintains a professional but friendly tone
        
        IMPORTANT RESTRICTIONS:
        - DO NOT include any follow-up offers or invitations for further assistance
        - DO NOT include phrases like "If you'd like, I can help...", "Just let me know...", "Feel free to ask...", or similar offers
        - Keep the response focused on the current query and suggestions only
        - End the response after providing the suggestions - do not add any closing offers
        
        Format the entire response as a single, valid HTML block following these requirements:
        - CRITICAL: Generate compact HTML without newline characters (\\n) between tags
        - Use <h2> or <h3> for section headings
        - Use <ul> and <li> for suggestions and lists
        - Use <p> tags for paragraphs
        - Use appropriate emojis: 📌 for suggestions, 💡 for tips, 🔍 for search advice
        - DO NOT include \\n characters anywhere - the HTML should be a single continuous string
        - Write in a clear, conversational style similar to ChatGPT
        
        Example structure:
        <h2>No Matching ADRs Found</h2>
        <p>I couldn't find any ADRs matching your query...</p>
        <h3>Suggestions</h3>
        <ul><li>Try...</li></ul>
        """
        
        prompt_template = ChatPromptTemplate.from_template(no_results_prompt)
        
        # Generate response using LLM
        response_chain = prompt_template | self.llm | StrOutputParser()
        response_text = response_chain.invoke({})
        
        # Clean up newline characters
        response_text = self._clean_html_response(response_text)
        
        # Return response with empty references array - no ADRs were found
        return {
            "query": query,
            "response": response_text,
            "references": []  # Explicitly empty - no documents found, so no references to show
        }

    def _get_dynamic_separators(self, text_content: str) -> List[str]:
        """
        Dynamically generates a list of separators based on recognized ADR section headings.
        This version uses a single, cleaner data structure to manage canonical headings and synonyms.

        Args:
            text_content: The full text content of the ADR document.

        Returns:
            A list of separators to be used by the text splitter.
        """
        found_separators = []

        for canonical_heading, synonyms in canonical_headings_map.items():
            # Create a regex pattern to find any of the synonyms for the current canonical heading.
            # This makes it case-insensitive and handles potential whitespace.
            pattern = r"^\s*(" + "|".join([re.escape(s) for s in synonyms]) + r")\s*$"

            if re.search(pattern, text_content, re.MULTILINE | re.IGNORECASE):
                # If any synonym is found, use the canonical heading to create the separator.
                separator = f"\n\n{canonical_heading}\n\n"
                if separator not in found_separators:
                    found_separators.append(separator)

        # Add general fallbacks to ensure chunks are created even if no headings are found.
        found_separators.extend(["\n\n", "\n", ". ", " "])

        return found_separators
