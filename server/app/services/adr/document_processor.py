"""Document processing service for extracting text from files."""
from fastapi import HTTPException, UploadFile
from typing import Tuple
import pypdfium2 as pdfium
import docx
from io import BytesIO

from app.utils.text_processing import clean_text_from_bullets


class DocumentProcessor:
    """Handles document processing and text extraction."""

    def extract_text_from_content(self, file_content: bytes, file_name: str) -> str:
        """
        Extract text from file content based on file extension.

        Args:
            file_content: The file content as bytes
            file_name: The filename

        Returns:
            Extracted text content

        Raises:
            HTTPException: If file format is unsupported or extraction fails
        """
        if not file_name:
            raise HTTPException(status_code=400, detail="File must have a filename")

        try:
            text_content = clean_text_from_bullets(
                self._extract_text_from_file_content(file_content, file_name)
            )
            return text_content
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Failed to extract text from document: {e}"
            )

    def _extract_text_from_pdf(self, file_obj: BytesIO) -> str:
        """Extract text from PDF using pypdfium2 from an in-memory object."""
        try:
            # pypdfium2's PdfDocument can directly accept a file-like object (BytesIO)
            pdf = pdfium.PdfDocument(file_obj)
            text_content = []

            for page_num in range(len(pdf)):
                page = pdf[page_num]
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                text_content.append(text)

                # Clean up
                textpage.close()
                page.close()

            pdf.close()
            return "\n".join(text_content)

        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")

    def _extract_text_from_docx(self, file_obj: BytesIO) -> str:
        """Extract text from Word document from an in-memory object."""
        try:
            # python-docx's Document can also accept a file-like object
            doc = docx.Document(file_obj)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            return "\n".join(text_content)

        except Exception as e:
            raise Exception(f"Failed to extract Word document text: {str(e)}")

    def _extract_text_from_file_content(
        self, file_content: bytes, filename: str
    ) -> str:
        """Extract text based on file extension from in-memory content."""
        file_ext = filename.lower().split(".")[-1]

        # Create an in-memory file object from the binary content.
        # This is a key step to make file-path-based libraries work.
        file_obj = BytesIO(file_content)

        if file_ext == "pdf":
            return self._extract_text_from_pdf(file_obj)
        elif file_ext in ["docx", "doc"]:
            return self._extract_text_from_docx(file_obj)
        elif file_ext in ["txt", "md"]:
            # For text files, we can just decode the bytes directly.
            return file_content.decode("utf-8")
        else:
            raise Exception(f"Unsupported file format: {file_ext}")

